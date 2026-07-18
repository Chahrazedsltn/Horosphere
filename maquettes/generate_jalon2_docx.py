#!/usr/bin/env python3
"""Génère le livrable Jalon 2 — Méthodologie & Conception UI/UX (docx avec images embarquées)."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Paths ──
BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = "/tmp"  # maquettes copies with simple names
OUT = os.path.join(BASE, "jalon2_methodologie_uiux_livrable.docx")

# Map: screenshot number -> (label, caption)
SCREENSHOTS = {
    1:  ("Écran 1 — Connexion", "Page de connexion Horosphere"),
    2:  ("Écran 2 — Tableau de bord Agent", "Dashboard agent avec pointage du jour et carte géofencing"),
    3:  ("Écran 3 — Historique des Présences", "Vue calendrier mensuel et détail des pointages"),
    4:  ("Écran 4 — Mon Profil", "Profil utilisateur, sécurité et préférences"),
    5:  ("Écran 5 — Vue RH — Tableau de bord", "Dashboard RH avec KPIs et liste employés"),
    6:  ("Écran 6 — Validation RH", "Anomalies & corrections avec filtres par type"),
    7:  ("Écran 7 — Paramètres Système", "Configuration géofencing, règles horaires, notifications"),
    8:  ("Écran 8 — Mes Demandes", "Formulaire nouvelle demande et historique"),
    9:  ("Écran 9 — Mes Documents", "Bulletins de paie, attestations et contrats"),
    10: ("Écran 10 — Gestion des Employés", "Liste paginée avec filtres et actions"),
    11: ("Écran 11 — Zones Géofencing", "Carte des zones, fiches sites et alertes temps réel"),
    12: ("Écran 12 — Rapports & Exports", "Exports mensuels et rapports personnalisés"),
}

# ── Styles helpers ──
ACCENT = RGBColor(0x25, 0x63, 0xEB)  # #2563eb
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level <= 2 else DARK
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
        set_cell_shading(cell, '2563EB')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F1F5F9')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_image(doc, img_num, width=6.2):
    path = os.path.join(IMG_DIR, f"maq_{img_num}.png")
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Image manquante : maq_{img_num}.png]", italic=True, color=RGBColor(0xDC, 0x26, 0x26))


def add_caption(doc, text):
    add_para(doc, text, italic=True, size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


# ══════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════
doc = Document()

# -- Default font --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ═══════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "JALON 2", bold=True, size=28, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Méthodologie & Conception UI/UX", bold=True, size=18, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
doc.add_paragraph()
add_para(doc, "HOROSPHERE", bold=True, size=22, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Application de Pointage Numérique", size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
doc.add_paragraph()
doc.add_paragraph()

add_table(doc, ["", ""], [
    ["Auteur", "Chahrazed Soltani"],
    ["Formation", "Bachelor CDA — IPSSI"],
    ["Session", "2025 / 2026"],
    ["Version", "Juillet 2026"],
], col_widths=[4, 8])

doc.add_page_break()

# ═══════════════════════════════════════════
#  SOMMAIRE
# ═══════════════════════════════════════════
add_heading(doc, "Sommaire", level=1)
sommaire = [
    "1.  Méthodologie de gestion de projet",
    "    1.1  Méthode choisie : Kanban adapté",
    "    1.2  Découpage temporel — Jalons",
    "    1.3  Backlog — User Stories priorisées",
    "    1.4  Outils de gestion",
    "2.  Charte graphique",
    "    2.1  Palette de couleurs",
    "    2.2  Typographies",
    "    2.3  Composants récurrents",
    "    2.4  Iconographie",
    "    2.5  Responsive & Layout",
    "3.  Parcours utilisateurs (User Flows)",
    "    3.1  Pointage d'une journée complète (Agent)",
    "    3.2  Traitement des anomalies (RH)",
    "    3.3  Soumission d'une demande (Agent)",
    "    3.4  Configuration du géofencing (Admin)",
    "4.  Maquettes et correspondance avec le CDCF",
    "5.  Principes UX appliqués",
    "6.  Matrice de traçabilité Écrans × CDCF",
]
for item in sommaire:
    add_para(doc, item, size=11, space_after=3)

doc.add_page_break()

# ═══════════════════════════════════════════
#  1. MÉTHODOLOGIE
# ═══════════════════════════════════════════
add_heading(doc, "1. Méthodologie de gestion de projet", level=1)

add_heading(doc, "1.1 Méthode choisie : Agile — Kanban adapté", level=2)
add_para(doc, "Le projet Horosphere est développé en solo dans le cadre d'une formation CDA. La méthode Kanban a été retenue pour sa flexibilité et son adaptation au travail individuel, contrairement à Scrum qui suppose une équipe pluridisciplinaire et des cérémonies collectives.")
add_para(doc, "Pourquoi Kanban plutôt que Scrum :", bold=True, space_after=3)
add_para(doc, "• Pas de rôles distincts (pas de Scrum Master ni de Product Owner séparés) — je cumule tous les rôles.", space_after=2)
add_para(doc, "• Le flux de travail est continu : les jalons mensuels du référentiel CDA structurent naturellement les itérations.", space_after=2)
add_para(doc, "• Le tableau Kanban (To Do → In Progress → Done) est suffisant pour suivre l'avancement sans cérémonie superflue.")

add_heading(doc, "1.2 Découpage temporel — Jalons", level=2)
add_para(doc, "Le projet est structuré en 6 jalons mensuels, alignés sur le calendrier du référentiel CDA :")
add_table(doc, ["Jalon", "Période", "Phase", "Livrable"], [
    ["J1", "Janvier 2026", "Cadrage & Rédaction", "Cahier des charges fonctionnel"],
    ["J2", "Février 2026", "Design UI/UX", "Maquettes Figma + méthodologie"],
    ["J3", "Mars 2026", "Modélisation BDD", "Schéma Merise (MCD/MLD/MPD)"],
    ["J4", "Avril 2026", "Architecture API", "Diagrammes UML + conception technique"],
    ["J5", "Mai 2026", "Développement & Tests", "Version bêta + suite PHPUnit"],
    ["J6", "Juin 2026", "Déploiement & Livraison", "Docker + documentation finale"],
], col_widths=[1.5, 3, 3.5, 5])

add_heading(doc, "1.3 Backlog — User Stories priorisées", level=2)
add_para(doc, "Le backlog est organisé par priorité (méthode MoSCoW) et par espace utilisateur.")

add_para(doc, "Must Have (indispensable)", bold=True, color=ACCENT, space_after=4)
add_table(doc, ["ID", "User Story", "CDCF"], [
    ["US01", "En tant qu'utilisateur, je veux me connecter avec mon email et mot de passe pour accéder à mon espace", "F1"],
    ["US02", "En tant qu'agent, je veux pointer mon arrivée en un clic pour enregistrer ma présence", "F2"],
    ["US03", "En tant qu'agent, je veux pointer mon départ pour clôturer ma journée", "F2"],
    ["US04", "En tant qu'agent, je veux mettre mon pointage en pause et le reprendre", "F2"],
    ["US05", "En tant que système, je veux vérifier la position GPS de l'agent par rapport au site", "F3"],
    ["US06", "En tant qu'agent, je veux consulter l'historique mensuel de mes pointages", "F4"],
    ["US07", "En tant que RH, je veux valider ou corriger les feuilles de temps", "F5"],
    ["US08", "En tant que RH, je veux recevoir des alertes automatiques sur les anomalies", "F6"],
    ["US09", "En tant que RH, je veux exporter les données de pointage en CSV ou PDF", "F7"],
    ["US10", "En tant qu'admin, je veux créer, modifier et désactiver des comptes", "F8"],
], col_widths=[1.5, 10, 1.5])

add_para(doc, "Should Have (important)", bold=True, color=ACCENT, space_after=4)
add_table(doc, ["ID", "User Story", "CDCF"], [
    ["US11", "En tant qu'admin, je veux configurer les sites avec coordonnées GPS et rayon de géofencing", "F9"],
    ["US12", "En tant qu'admin, je veux consulter les logs d'activité pour la supervision technique", "F10"],
    ["US13", "En tant qu'agent, je veux soumettre une demande de congé ou de correction", "F5"],
    ["US14", "En tant qu'agent, je veux consulter et télécharger mes documents", "F4"],
    ["US15", "En tant qu'utilisateur, je veux réinitialiser mon mot de passe oublié", "F1"],
], col_widths=[1.5, 10, 1.5])

add_para(doc, "Could Have (souhaitable)", bold=True, color=ACCENT, space_after=4)
add_table(doc, ["ID", "User Story", "CDCF"], [
    ["US16", "En tant qu'agent, je veux personnaliser mes préférences de notification", "—"],
    ["US17", "En tant que RH, je veux voir un graphique du taux de présence mensuel", "F7"],
    ["US18", "En tant qu'admin, je veux choisir un thème visuel pour l'interface", "—"],
], col_widths=[1.5, 10, 1.5])

add_heading(doc, "1.4 Outils de gestion", level=2)
add_para(doc, "• Tableau Kanban : GitHub Projects (colonnes To Do / In Progress / Review / Done)", space_after=2)
add_para(doc, "• Versioning : Git + GitHub (branches feature/*)", space_after=2)
add_para(doc, "• CI/CD : GitHub Actions (lint, tests PHPUnit)", space_after=2)
add_para(doc, "• Maquettage : Figma")

doc.add_page_break()

# ═══════════════════════════════════════════
#  2. CHARTE GRAPHIQUE
# ═══════════════════════════════════════════
add_heading(doc, "2. Charte graphique", level=1)
add_para(doc, "L'application utilise un système de design à thèmes interchangeables avec une base constante. Le thème par défaut est Slate-Blue (Ardoise + Bleu Électrique).")

add_heading(doc, "2.1 Palette de couleurs", level=2)

add_para(doc, "Couleurs de contenu (constantes sur tous les thèmes)", bold=True, space_after=4)
add_table(doc, ["Rôle", "Variable CSS", "Hex", "Usage"], [
    ["Fond principal", "--bg", "#f8fafc", "Arrière-plan général de l'application"],
    ["Surface carte", "--surface", "#ffffff", "Cartes, modales, formulaires"],
    ["Surface secondaire", "--surface2", "#f1f5f9", "Lignes alternées, zones secondaires"],
    ["Bordure", "--border", "#e2e8f0", "Séparateurs, contours de cartes"],
    ["Texte principal", "--text", "#0f172a", "Titres, contenus importants"],
    ["Texte secondaire", "--text2", "#475569", "Labels, descriptions"],
    ["Texte tertiaire", "--text3", "#94a3b8", "Placeholders, métadonnées"],
], col_widths=[3, 3, 2.5, 5])

add_para(doc, "Couleurs d'accentuation (thème Slate-Blue — défaut)", bold=True, space_after=4)
add_table(doc, ["Rôle", "Variable CSS", "Hex", "Usage"], [
    ["Accent principal", "--accent", "#2563eb", "Boutons primaires, liens actifs, sidebar active"],
    ["Accent clair", "--accent-light", "#eff6ff", "Fond des badges, survol léger"],
    ["Accent moyen", "--accent-mid", "#3b82f6", "Hover des boutons"],
    ["Sidebar", "--sidebar-bg", "#0f172a", "Fond de la barre latérale (ardoise foncé)"],
], col_widths=[3, 3, 2.5, 5])

add_para(doc, "Couleurs de statut (constantes)", bold=True, space_after=4)
add_table(doc, ["Rôle", "Couleur", "Hex", "Usage"], [
    ["Succès", "Vert", "#16a34a", "Validé, Présent, Approuvé"],
    ["Fond succès", "Vert clair", "#f0fdf4", "Badge vert (fond)"],
    ["Erreur", "Rouge", "#dc2626", "Anomalie, Absent, Rejeté"],
    ["Fond erreur", "Rouge clair", "#fef2f2", "Badge rouge (fond)"],
    ["Avertissement", "Ambre", "#ca8a04", "En attente, En pause, Alerte"],
    ["Fond avertissement", "Ambre clair", "#fefce8", "Badge ambre (fond)"],
], col_widths=[3, 3, 2.5, 5])

add_para(doc, "Justification des choix de couleurs :", bold=True, space_after=4)
add_para(doc, "Le fond clair (#f8fafc) avec sidebar sombre crée un contraste fort qui guide l'œil vers le contenu principal. Ce pattern est largement adopté dans les applications SaaS professionnelles (Slack, Linear, Notion). Les couleurs de statut suivent les conventions sémantiques universelles (vert = OK, rouge = erreur, ambre = attention), ce qui réduit la charge cognitive pour les utilisateurs non techniques comme Marie (agent de terrain). Le bleu électrique comme accent est perçu comme professionnel et fiable, adapté à un outil RH.")

add_heading(doc, "2.2 Typographies", level=2)
add_table(doc, ["Usage", "Police", "Justification"], [
    ["Interface générale", "Outfit (sans-serif)", "Police géométrique moderne, excellente lisibilité sur écran, bonne couverture de graisses (300–700). Son dessin ouvert et ses formes rondes transmettent un sentiment d'accessibilité."],
    ["Données numériques", "JetBrains Mono (monospace)", "Alignement parfait des chiffres dans les tableaux de pointage (heures, durées). Les chiffres sont à largeur fixe, ce qui facilite la lecture comparative des colonnes."],
], col_widths=[3, 4, 7])

add_heading(doc, "2.3 Composants récurrents", level=2)
add_table(doc, ["Composant", "Style", "Usage"], [
    ["Cartes", "Fond blanc, border-radius 12px, ombre légère", "Conteneurs principaux (pointage, statistiques, formulaires)"],
    ["Boutons primaires", "Fond accent, texte blanc, radius 12px", "Actions principales (Se connecter, Pointer, Valider)"],
    ["Badges de statut", "Texte coloré sur fond clair, radius pill", "Affichage des états (Validé, En attente, Absent)"],
    ["Sidebar", "Fond ardoise, items icon + label, actif en accent", "Navigation principale sur toutes les pages"],
    ["Tableaux", "Headers gris clair, lignes alternées", "Historique, liste employés, détail pointages"],
    ["Indicateurs KPI", "Nombre en grand (font-mono), label, carte", "Dashboards agent et RH"],
], col_widths=[3, 5, 5.5])

add_heading(doc, "2.4 Iconographie", level=2)
add_para(doc, "Les icônes utilisent la bibliothèque Lucide React (fork de Feather Icons), choisie pour son style linéaire minimaliste cohérent avec le design épuré, sa légèreté (tree-shakable) et sa couverture complète des cas d'usage métier (horloge, carte, utilisateur, document, alerte).")

add_heading(doc, "2.5 Responsive & Layout", level=2)
add_table(doc, ["Dimension", "Valeur", "Justification"], [
    ["Largeur sidebar", "225px", "Suffisant pour label + icône sans empiéter sur le contenu"],
    ["Hauteur topbar", "60px", "Horloge temps réel et profil toujours visibles"],
    ["Border-radius", "12px", "Arrondi moderne cohérent sur tous les composants"],
    ["Breakpoint mobile", "< 768px", "Sidebar rétractée, layout en colonne unique"],
    ["Taille min boutons", "48px hauteur", "Cible tactile confortable (recommandation WCAG)"],
], col_widths=[3, 3, 7.5])
add_para(doc, "L'application est conçue mobile-first conformément au CDCF : l'agent de terrain accède depuis son smartphone. Les boutons de pointage sont dimensionnés pour une utilisation tactile.")

doc.add_page_break()

# ═══════════════════════════════════════════
#  3. PARCOURS UTILISATEURS
# ═══════════════════════════════════════════
add_heading(doc, "3. Parcours utilisateurs (User Flows)", level=1)

add_heading(doc, "3.1 Flux principal — Pointage d'une journée complète (Agent)", level=2)
add_para(doc, "Ce parcours couvre les fonctionnalités F1, F2, F3 du CDCF.", italic=True, color=GRAY)

steps = [
    ("1.", "L'agent ouvre l'application et saisit ses identifiants (email / mot de passe)."),
    ("2.", "Après connexion, il arrive sur le tableau de bord affichant le statut « Non pointé »."),
    ("3.", "Il clique sur le bouton vert Arrivée (action principale, la plus visible)."),
    ("4.", "Le navigateur demande l'autorisation de géolocalisation. Les coordonnées GPS sont capturées."),
    ("5.", "Le système compare la position avec le centre du site assigné (formule de Haversine)."),
    ("6.", "Si la distance < rayon (200m) : statut EN_COURS, zone validée."),
    ("7.", "Si la distance > rayon : statut HORS_ZONE, alerte envoyée au RH, pointage NON bloqué (exigence juridique F3)."),
    ("8.", "En journée, l'agent peut cliquer Pause (statut EN_PAUSE, compteur gelé) puis Reprendre."),
    ("9.", "En fin de journée, il clique Départ : statut VALIDE, durée totale affichée."),
]
for num, text in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_num = p.add_run(num + " ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = ACCENT
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, "3.2 Flux RH — Traitement des anomalies", level=2)
add_para(doc, "Ce parcours couvre les fonctionnalités F5, F6, F7 du CDCF.", italic=True, color=GRAY)
steps_rh = [
    ("1.", "Le RH se connecte et accède au dashboard RH. Un badge « 3 anomalies à traiter » est visible."),
    ("2.", "Il clique sur « Validation RH » dans la sidebar."),
    ("3.", "La page affiche les anomalies filtrables par type : Retards, Oublis de pointage, Hors zone."),
    ("4.", "Pour chaque anomalie, le RH voit le détail contextuel (heure, écart, lieu, motif)."),
    ("5.", "Il choisit : Approuver (valide le pointage tel quel), Corriger la saisie (saisie manuelle), ou Refuser."),
    ("6.", "Chaque action est journalisée dans l'audit log avec l'identité du RH et l'horodatage."),
]
for num, text in steps_rh:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_num = p.add_run(num + " ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = ACCENT
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, "3.3 Flux Agent — Soumission d'une demande de congé", level=2)
add_para(doc, "Ce parcours couvre la fonctionnalité F5 du CDCF.", italic=True, color=GRAY)
steps_dem = [
    ("1.", "L'agent clique sur « Mes Demandes » dans la sidebar."),
    ("2.", "Il clique sur « + Nouvelle demande »."),
    ("3.", "Il remplit le formulaire : type (Congé), dates début/fin, motif optionnel, justificatif optionnel."),
    ("4.", "Il clique « Envoyer ». La demande apparaît dans l'historique avec le statut « En attente »."),
    ("5.", "Le RH reçoit une notification et peut approuver ou rejeter la demande."),
]
for num, text in steps_dem:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_num = p.add_run(num + " ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = ACCENT
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, "3.4 Flux Admin — Configuration du géofencing", level=2)
add_para(doc, "Ce parcours couvre la fonctionnalité F9 du CDCF.", italic=True, color=GRAY)
steps_geo = [
    ("1.", "L'admin clique sur « Zones Géo » dans la sidebar."),
    ("2.", "La carte interactive affiche les zones existantes (cercles)."),
    ("3.", "Il clique « + Nouvelle zone » et saisit l'adresse, le rayon, le statut (Actif/Test)."),
    ("4.", "Le geocoding automatique place le marqueur sur la carte."),
    ("5.", "Il sauvegarde. La zone apparaît sur la carte et les alertes géofencing sont actives."),
]
for num, text in steps_geo:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_num = p.add_run(num + " ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = ACCENT
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════
#  4. MAQUETTES ET CORRESPONDANCE CDCF
# ═══════════════════════════════════════════
add_heading(doc, "4. Maquettes et correspondance avec le CDCF", level=1)
add_para(doc, "Chaque écran de l'application répond à un ou plusieurs besoins fonctionnels identifiés dans le cahier des charges (F1–F10). Les captures ci-dessous sont issues des maquettes Figma.")

# Screen descriptions with CDCF mapping
screen_details = {
    1: {
        "fonc": "F1 : Authentification",
        "desc": "Formulaire email + mot de passe, centré et minimaliste. Lien « Mot de passe oublié ? » pour la récupération. Option « Rester connecté » pour la persistance de session. Logo Horosphere pour identifier l'application.",
        "cdcf": "Répond à F1 — « Connexion sécurisée par email/mot de passe avec récupération de mot de passe ».",
    },
    2: {
        "fonc": "F2 : Pointage — F3 : Géolocalisation — F6 : Alertes",
        "desc": "Zone gauche : trois boutons d'action Arrivée (vert), Pause (ambre), Départ (rouge) avec indicateurs temps réel (durée travaillée, statut). Zone droite : carte Google Maps avec zone de géofencing, indicateurs « Zone valide » et distance. Zone basse : 4 KPIs mensuels + tableau des derniers pointages.",
        "cdcf": "F2 — Boutons Arrivée/Pause/Départ en 2 clics. F3 — Carte géofencing avec vérification de position. F6 — Statut anomalie visible directement.",
    },
    3: {
        "fonc": "F4 : Historique",
        "desc": "Vue calendrier mensuel avec code couleur (bleu = présent, rose = absence). Navigation mois par mois. Tableau de détail à droite avec jour, date, entrée, sortie, total. Bouton Export CSV. Total mensuel affiché.",
        "cdcf": "Répond à F4 — « Consultation mensuelle des heures et sites en vue liste ou calendrier ».",
    },
    4: {
        "fonc": "F1 : Authentification (gestion du compte)",
        "desc": "Informations personnelles modifiables. Section sécurité pour le changement de mot de passe. Préférences de notification (email, SMS, résumé hebdomadaire).",
        "cdcf": "Complète F1 — gestion du compte et sécurité.",
    },
    5: {
        "fonc": "F5 : Validation — F6 : Alertes — F7 : Export",
        "desc": "5 KPIs en bannière (employés actifs, présents, absents, pointages en attente, congés). Graphique hebdomadaire des présences. Taux de présence et ponctualité. Alerte « 3 employés absents sans justificatif ». Liste des employés avec statut du jour et actions.",
        "cdcf": "F5 — Vue globale pour la validation. F6 — Alertes anomalies dans le dashboard. F7 — Bouton Exporter dans la liste.",
    },
    6: {
        "fonc": "F5 : Validation — F6 : Alertes",
        "desc": "Filtres par onglets : Tous, Retards, Oublis de pointage, Hors zone. Fiche par anomalie avec détail contextuel (heure, écart, lieu, motif). Actions : Approuver / Corriger / Refuser. Les 3 types d'alertes du CDCF sont représentés.",
        "cdcf": "Répond à F5 — « Approbation des feuilles de temps et traitement des anomalies ». F6 — Retard, oubli de sortie, hors zone.",
    },
    7: {
        "fonc": "F9 : Configuration des sites",
        "desc": "Configuration géofencing : adresse, rayon, carte, ajout de sites. Règles horaires : début/fin journée, seuil retard, durée max pause. Notifications & automatisations avec toggles.",
        "cdcf": "Répond à F9 — « Création de fiches clients avec coordonnées GPS et paramétrage du rayon de géofencing ».",
    },
    8: {
        "fonc": "F5 : Validation (soumission côté agent)",
        "desc": "Formulaire inline : type de demande, dates, motif. Historique des demandes avec statuts colorés (En attente, Approuvée, Refusée). Bouton Annuler disponible tant que le statut est En attente.",
        "cdcf": "Complète F5 côté agent — soumission des demandes de congé et de correction.",
    },
    9: {
        "fonc": "F7 : Export (consultation côté agent)",
        "desc": "Bulletins de paie en cards avec aperçu et téléchargement. Attestations & contrats en tableau. Filtres par catégorie et barre de recherche.",
        "cdcf": "Complète F7 — les documents générés (CSV/PDF) sont accessibles par l'agent.",
    },
    10: {
        "fonc": "F8 : Gestion des utilisateurs",
        "desc": "KPIs : total employés, présents, nouveaux, désactivés. Filtres : recherche, département, statut. Tableau paginé avec pointage du jour, lieu, responsable. Actions : Éditer, Voir, Alerte. Boutons Ajouter et Exporter.",
        "cdcf": "Répond à F8 — « Création, modification et désactivation des comptes ».",
    },
    11: {
        "fonc": "F9 : Configuration des sites — F3 : Géolocalisation",
        "desc": "Carte interactive avec zones de géofencing visualisées. Fiches site : adresse, rayon, employés assignés, présents. Alertes géofencing temps réel en tableau. Bouton Nouvelle zone.",
        "cdcf": "F9 — Configuration des sites. F3 — Alertes de géofencing en temps réel.",
    },
    12: {
        "fonc": "F7 : Export",
        "desc": "Export mensuel : mois, format (CSV/XLSX/PDF), filtre employé. Rapport personnalisé : dates, employé(s), type. Graphique annuel du taux de présence.",
        "cdcf": "Répond à F7 — « Génération de fichiers CSV/PDF pour la paie et la preuve de service client ».",
    },
}

for num in range(1, 13):
    label, caption = SCREENSHOTS[num]
    details = screen_details[num]

    add_heading(doc, label, level=2)
    add_para(doc, f"→ {details['fonc']}", bold=True, color=ACCENT, space_after=6)

    add_image(doc, num)
    add_caption(doc, caption)

    add_para(doc, details["desc"])
    add_para(doc, f"Lien CDCF : {details['cdcf']}", bold=True, size=10, space_after=12)

    if num < 12:
        doc.add_page_break()

doc.add_page_break()

# ═══════════════════════════════════════════
#  5. PRINCIPES UX
# ═══════════════════════════════════════════
add_heading(doc, "5. Principes UX appliqués", level=1)

add_heading(doc, "5.1 Loi de Fitts — Taille et positionnement des cibles", level=2)
add_para(doc, "Les boutons de pointage (Arrivée, Pause, Départ) sont les plus grands éléments interactifs de l'écran principal. Leur taille (pleine largeur de la carte) et leur position centrale réduisent le temps d'acquisition conformément à la loi de Fitts. Cela répond directement au besoin de Marie (agent) : « pointer en deux clics, sans formation technique ».")

add_heading(doc, "5.2 Hiérarchie visuelle — Couleurs sémantiques", level=2)
add_para(doc, "Les couleurs de statut (vert/rouge/ambre) sont utilisées de manière cohérente sur tous les écrans : vert = action positive (Validé, Présent, Arrivée), rouge = action critique (Anomalie, Absent, Départ), ambre = état intermédiaire (En pause, En attente). Cette cohérence réduit la charge cognitive et permet une lecture instantanée des tableaux de bord.")

add_heading(doc, "5.3 Progressive disclosure — Information à la demande", level=2)
add_para(doc, "Le dashboard agent affiche uniquement les informations essentielles (pointage du jour, KPIs). Les détails (historique, documents, demandes) sont accessibles via la navigation latérale. Le RH voit les KPIs de synthèse en premier, avec un badge d'alerte « 3 anomalies » qui invite à explorer la page de validation.")

add_heading(doc, "5.4 Consistance — Navigation unifiée", level=2)
add_para(doc, "La sidebar est identique en structure pour tous les profils, seuls les items changent. Agent : Accueil, Mon Historique, Mes Demandes, Mes Documents. RH/Admin : Vue Globale, Employés, Validation RH, Rapports + section Configuration. L'utilisateur en bas de la sidebar (avatar + nom + rôle) est toujours visible, renforçant le sentiment de contexte.")

doc.add_page_break()

# ═══════════════════════════════════════════
#  6. MATRICE DE TRAÇABILITÉ
# ═══════════════════════════════════════════
add_heading(doc, "6. Matrice de traçabilité — Écrans × Fonctionnalités CDCF", level=1)
add_para(doc, "Le tableau ci-dessous démontre que chaque fonctionnalité du CDCF est couverte par au moins un écran de l'application.")

add_table(doc, ["Fonctionnalité CDCF", "Écrans correspondants"], [
    ["F1 — Authentification", "Écran 1 (Connexion), Écran 4 (Mon Profil)"],
    ["F2 — Pointage", "Écran 2 (Dashboard Agent)"],
    ["F3 — Géolocalisation & Geofencing", "Écran 2 (carte géofencing), Écran 11 (Zones Géo)"],
    ["F4 — Historique", "Écran 3 (Historique — calendrier + détail)"],
    ["F5 — Validation", "Écran 6 (Validation RH), Écran 8 (Mes Demandes)"],
    ["F6 — Alertes", "Écrans 2, 5, 6 (anomalies visibles à chaque niveau)"],
    ["F7 — Export", "Écran 9 (Mes Documents), Écran 12 (Rapports & Exports)"],
    ["F8 — Gestion utilisateurs", "Écran 10 (Gestion des Employés)"],
    ["F9 — Configuration des sites", "Écran 7 (Paramètres), Écran 11 (Zones Géofencing)"],
    ["F10 — Supervision technique", "Écran 5 (Dashboard RH), Écran 7 (Paramètres)"],
], col_widths=[5, 9])

add_para(doc, "Couverture : 10/10 fonctionnalités du CDCF sont couvertes par au moins un écran.", bold=True, color=ACCENT, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
doc.save(OUT)
print(f"Document généré : {OUT}")
print(f"Taille : {os.path.getsize(OUT) / 1024:.0f} Ko")
